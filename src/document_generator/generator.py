"""
Main document generator for DPDP-compliant legal templates.
Generates Word documents pre-filled with business profile data.
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
from pathlib import Path
import secrets
from typing import Dict, List, Tuple, Optional
import zipfile

from .constants import (
    DATA_TYPE_DISPLAY_NAMES,
    DATA_TYPE_DETAILED_DESCRIPTIONS,
    RETENTION_REQUIREMENTS,
    LEGAL_BASIS_MAP,
    THIRD_SCHEDULE_THRESHOLDS
)
from .validators import validate_profile, sanitize_input, ValidationError

class DocumentGenerationError(Exception):
    """Raised when document generation fails"""
    pass

class DocumentGenerator:
    """
    Generate DPDP-compliant legal documents from business profile.
    
    IMPORTANT: Generated documents are TEMPLATES requiring legal review.
    
    Usage:
        generator = DocumentGenerator(business_profile, gap_analysis)
        documents = generator.generate_all_required_documents()
        zip_path = generator.export_all_to_zip(documents, 'output.zip')
    """
    
    def __init__(self, business_profile: dict, gap_analysis: dict):
        """
        Initialize document generator.
        
        Args:
            business_profile: Business profile from assessment questionnaire
            gap_analysis: Gap analysis results with compliance score
            
        Raises:
            ValidationError: If profile validation fails
            DocumentGenerationError: If templates directory not found
        """
        # Validate inputs
        validated_profile = validate_profile(business_profile)
        
        self.profile = validated_profile
        self.gaps = gap_analysis
        self.templates_dir = Path(__file__).parent / "templates"
        
        # Verify templates exist
        if not self.templates_dir.exists():
            raise DocumentGenerationError(
                f"Templates directory not found: {self.templates_dir}"
            )
    
    def generate_privacy_notice(self) -> Document:
        """
        Generate Privacy Notice per Rule 3.
        
        Required by: ALL businesses
        Penalty for non-compliance: Rs. 50 crore
        
        Returns:
            Document object with privacy notice
        """
        template_path = self.templates_dir / "privacy_notice.docx"
        if not template_path.exists():
            raise DocumentGenerationError(f"Template not found: {template_path}")
        
        doc = Document(str(template_path))
        
        replacements = {
            '{{BUSINESS_NAME}}': self._get_safe_value('business_name'),
            '{{ADDRESS}}': self._get_safe_value('address', '[INSERT BUSINESS ADDRESS]'),
            '{{EMAIL}}': self._get_safe_value('email', '[INSERT CONTACT EMAIL]'),
            '{{PHONE}}': self._get_safe_value('phone', '[INSERT CONTACT PHONE]'),
            '{{DATA_TYPES}}': self._format_data_types_readable(),
            '{{PURPOSES}}': self._format_purposes(),
            '{{DATE}}': datetime.now().strftime("%B %d, %Y"),
            '{{WITHDRAWAL_METHOD}}': self._format_withdrawal_method(),
            '{{GRIEVANCE_CONTACT}}': self._get_safe_value('email', '[INSERT GRIEVANCE EMAIL]')
        }
        
        self._replace_all_placeholders(doc, replacements)
        self._add_document_metadata(doc, 'Privacy Notice', 'Rule 3')
        self._add_template_disclaimer(doc)
        
        return doc
    
    def generate_consent_form(self) -> Document:
        """
        Generate Consent Form per Section 6.
        
        Required by: ALL businesses
        Penalty for non-compliance: Rs. 50 crore
        
        Returns:
            Document object with consent form
        """
        template_path = self.templates_dir / "consent_form.docx"
        if not template_path.exists():
            raise DocumentGenerationError(f"Template not found: {template_path}")
        
        doc = Document(str(template_path))
        
        replacements = {
            '{{BUSINESS_NAME}}': self._get_safe_value('business_name'),
            '{{DATA_COLLECTED}}': self._format_data_types_detailed(),
            '{{PURPOSES}}': self._format_purposes(),
            '{{RETENTION_PERIOD}}': self._format_retention_summary(),
            '{{WITHDRAWAL_METHOD}}': self._format_withdrawal_method(),
            '{{DATE}}': datetime.now().strftime("%B %d, %Y")
        }
        
        self._replace_all_placeholders(doc, replacements)
        self._add_document_metadata(doc, 'Consent Form', 'Section 6')
        self._add_template_disclaimer(doc)
        
        return doc
    
    def generate_grievance_procedure(self) -> Document:
        """
        Generate Grievance Redressal Procedure per Rule 14.
        
        Required by: ALL businesses
        Penalty for non-compliance: Rs. 50 crore
        
        Returns:
            Document object with grievance procedure
        """
        template_path = self.templates_dir / "grievance_procedure.docx"
        if not template_path.exists():
            raise DocumentGenerationError(f"Template not found: {template_path}")
        
        doc = Document(str(template_path))
        
        replacements = {
            '{{BUSINESS_NAME}}': self._get_safe_value('business_name'),
            '{{GRIEVANCE_EMAIL}}': self._get_safe_value('email', '[INSERT GRIEVANCE EMAIL]'),
            '{{GRIEVANCE_PHONE}}': self._get_safe_value('phone', '[INSERT GRIEVANCE PHONE]'),
            '{{DATE}}': datetime.now().strftime("%B %d, %Y")
        }
        
        self._replace_all_placeholders(doc, replacements)
        self._add_document_metadata(doc, 'Grievance Redressal Procedure', 'Rule 14')
        self._add_template_disclaimer(doc)
        
        return doc
    
    def generate_retention_schedule(self) -> Document:
        """
        Generate Data Retention Schedule per Rule 8.
        
        Required by: ALL businesses
        Penalty for non-compliance: Rs. 50 crore
        
        Returns:
            Document object with retention schedule
        """
        template_path = self.templates_dir / "retention_schedule.docx"
        if not template_path.exists():
            raise DocumentGenerationError(f"Template not found: {template_path}")
        
        doc = Document(str(template_path))
        
        # Populate retention table
        self._populate_retention_table(doc)
        
        replacements = {
            '{{BUSINESS_NAME}}': self._get_safe_value('business_name'),
            '{{DATE}}': datetime.now().strftime("%B %d, %Y")
        }
        self._replace_all_placeholders(doc, replacements)
        
        self._add_document_metadata(doc, 'Data Retention Schedule', 'Rule 8')
        self._add_template_disclaimer(doc)
        
        return doc
    
    def generate_breach_notification_templates(self) -> Tuple[Document, Document]:
        """
        Generate breach notification templates for crisis preparedness.
        
        Required by: ALL businesses (for future use during breach)
        Penalty for non-compliance: Rs. 200 crore
        
        Returns:
            Tuple of (DPB notification, Data Principal notification)
        """
        # Data Protection Board notification
        template_dpb = self.templates_dir / "breach_notification_dpb.docx"
        if not template_dpb.exists():
            raise DocumentGenerationError(f"Template not found: {template_dpb}")
        
        doc_dpb = Document(str(template_dpb))
        
        replacements_dpb = {
            '{{BUSINESS_NAME}}': self._get_safe_value('business_name'),
            '{{ADDRESS}}': self._get_safe_value('address', '[INSERT]'),
            '{{EMAIL}}': self._get_safe_value('email', '[INSERT]'),
            '{{PHONE}}': self._get_safe_value('phone', '[INSERT]'),
            '{{DATA_CATEGORIES}}': self._format_data_types_readable(),
            '{{BREACH_DATE}}': '[INSERT: Date breach discovered - DD/MM/YYYY]',
            '{{BREACH_NATURE}}': '[INSERT: Unauthorized access / Data leak / System compromise]',
            '{{AFFECTED_COUNT}}': '[INSERT: Number of users affected]',
            '{{CONSEQUENCES}}': '[INSERT: Identity theft risk / Financial fraud risk / Privacy violation]',
            '{{MITIGATION}}': '[INSERT: Password reset / Notification sent / System patched / Investigation ongoing]'
        }
        
        self._replace_all_placeholders(doc_dpb, replacements_dpb)
        self._add_breach_template_header(doc_dpb, 'Data Protection Board')
        self._add_document_metadata(doc_dpb, 'Breach Notification (DPB)', 'Rule 7')
        
        # Data Principal notification
        template_users = self.templates_dir / "breach_notification_user.docx"
        if not template_users.exists():
            raise DocumentGenerationError(f"Template not found: {template_users}")
        
        doc_users = Document(str(template_users))
        
        replacements_users = {
            '{{BUSINESS_NAME}}': self._get_safe_value('business_name'),
            '{{GRIEVANCE_CONTACT}}': self._get_safe_value('email', '[INSERT]'),
            '{{BREACH_DESCRIPTION}}': '[INSERT: Simple language - "Someone accessed our database without permission"]',
            '{{YOUR_DATA}}': '[INSERT: "Your email address and phone number may have been accessed"]',
            '{{WHAT_WE_DID}}': '[INSERT: "We immediately secured the system and reset all passwords"]',
            '{{WHAT_YOU_SHOULD_DO}}': '[INSERT: "Change your password immediately at [link]"]'
        }
        
        self._replace_all_placeholders(doc_users, replacements_users)
        self._add_breach_template_header(doc_users, 'Data Principals')
        self._add_document_metadata(doc_users, 'Breach Notification (Users)', 'Rule 7')
        
        return doc_dpb, doc_users
    
    def generate_parental_consent_form(self) -> Optional[Document]:
        """
        Generate Parental Consent Form per Rule 10.
        
        Required by: Businesses processing children's data (under 18)
        Penalty for non-compliance: Rs. 200 crore
        
        Returns:
            Document object or None if not applicable
        """
        if not self.profile.get('processes_children_data', False):
            return None
        
        template_path = self.templates_dir / "parental_consent.docx"
        if not template_path.exists():
            raise DocumentGenerationError(f"Template not found: {template_path}")
        
        doc = Document(str(template_path))
        
        replacements = {
            '{{BUSINESS_NAME}}': self._get_safe_value('business_name'),
            '{{SERVICE_DESCRIPTION}}': '[INSERT: Describe your service/app/platform]',
            '{{DATA_COLLECTED}}': self._format_data_types_readable(),
            '{{PURPOSES}}': self._format_purposes(),
            '{{VERIFICATION_METHOD}}': '[INSERT: ID verification / Payment method / Email confirmation]',
            '{{DATE}}': datetime.now().strftime("%B %d, %Y")
        }
        
        self._replace_all_placeholders(doc, replacements)
        self._add_document_metadata(doc, 'Parental Consent Form', 'Rule 10')
        self._add_template_disclaimer(doc)
        
        return doc
    
    def generate_processor_agreement_checklist(self) -> Optional[Document]:
        """
        Generate Data Processor Agreement checklist per Rule 6(1)(f).
        
        Required by: Businesses using data processors (vendors)
        Penalty for non-compliance: Rs. 250 crore
        
        NOTE: Generates CHECKLIST only, not full legal agreement.
        
        Returns:
            Document object or None if not applicable
        """
        if not self.profile.get('has_processors', False):
            return None
        
        doc = Document()
        
        # Title
        doc.add_heading('Data Processor Agreement - Required Clauses Checklist', 0)
        
        # Warning
        warning = doc.add_paragraph()
        warning.add_run('CRITICAL WARNING: ').bold = True
        warning.add_run(
            'This is NOT a legal agreement. This is a checklist of clauses required by '
            'DPDP Act Rule 6(1)(f). Have your lawyer draft the actual Data Processor '
            'Agreement including all items below.\n\n'
            'Penalty for non-compliant DPA: Rs. 250 crore (security breach)'
        )
        
        # Required clauses
        required_clauses = [
            ('Security Safeguards', 
             'Processor must implement encryption, access control, logging, and backups per Rule 6.'),
            
            ('Processing Instructions', 
             'Processor may only process data per your documented written instructions.'),
            
            ('Sub-Processor Authorization', 
             'Processor must obtain your prior written consent before engaging sub-processors.'),
            
            ('Data Principal Rights Assistance', 
             'Processor must assist you in responding to Data Principal access, correction, erasure requests.'),
            
            ('Breach Notification', 
             'Processor must notify you within 24 hours of discovering any data breach.'),
            
            ('Audit Rights', 
             'You have right to audit processor compliance annually or upon reasonable notice.'),
            
            ('Data Return/Deletion', 
             'Upon termination, processor must return or certify deletion of all data within 30 days.'),
            
            ('Personnel Confidentiality', 
             'All processor personnel must be bound by confidentiality obligations.'),
            
            ('Liability for Failures', 
             'Processor liable for damages from security failures or unauthorized disclosure.'),
            
            ('Indemnification', 
             'Processor indemnifies you for penalties arising from processor DPDP violations.')
        ]
        
        for clause_name, description in required_clauses:
            doc.add_heading(f'[ ] {clause_name}', level=2)
            doc.add_paragraph(description)
            
            lawyer_note = doc.add_paragraph()
            lawyer_note.add_run('YOUR LAWYER MUST DRAFT: ').bold = True
            lawyer_note.add_run(
                'Specific legal language, breach remedies, liability caps, and termination rights.'
            )
            doc.add_paragraph('')
        
        self._add_document_metadata(doc, 'DPA Checklist', 'Rule 6(1)(f)')
        
        return doc
    
    def generate_all_required_documents(self) -> Dict[str, Document]:
        """
        Generate all documents required based on profile and gaps.
        
        Returns:
            Dictionary mapping document names to Document objects
        """
        documents = {}
        
        # Universal documents
        documents['01_Privacy_Notice'] = self.generate_privacy_notice()
        documents['02_Consent_Form'] = self.generate_consent_form()
        documents['03_Grievance_Procedure'] = self.generate_grievance_procedure()
        documents['04_Retention_Schedule'] = self.generate_retention_schedule()
        
        # Breach templates (pre-crisis preparation)
        breach_dpb, breach_users = self.generate_breach_notification_templates()
        documents['05_Breach_Notification_DPB'] = breach_dpb
        documents['06_Breach_Notification_Users'] = breach_users
        
        # Conditional documents
        if self.profile.get('processes_children_data'):
            parental_consent = self.generate_parental_consent_form()
            if parental_consent:
                documents['07_Parental_Consent_Form'] = parental_consent
        
        if self.profile.get('has_processors'):
            processor_checklist = self.generate_processor_agreement_checklist()
            if processor_checklist:
                documents['08_Processor_Agreement_Checklist'] = processor_checklist
        
        return documents
    
    def export_to_docx(self, document: Document, filename: str) -> Path:
        """
        Export document to DOCX file.
        
        Args:
            document: Document object to export
            filename: Output filename
            
        Returns:
            Path to exported file
        """
        output_dir = Path("data/processed/generated_documents")
        output_dir.mkdir(parents=True, exist_ok=True)
        
        filepath = output_dir / filename
        document.save(str(filepath))
        
        return filepath
    
    def export_all_to_zip(self, documents: Dict[str, Document], zip_filename: str) -> Path:
        """
        Export all documents to ZIP file.
        
        Args:
            documents: Dictionary of document name to Document object
            zip_filename: Output ZIP filename
            
        Returns:
            Path to ZIP file
        """
        output_dir = Path("data/processed/generated_documents")
        output_dir.mkdir(parents=True, exist_ok=True)
        
        zip_path = output_dir / zip_filename
        
        with zipfile.ZipFile(zip_path, 'w', zipfile.ZIP_DEFLATED) as zipf:
            for doc_name, doc in documents.items():
                temp_path = self.export_to_docx(doc, f"{doc_name}.docx")
                zipf.write(temp_path, f"{doc_name}.docx")
        
        return zip_path
    
    # ========================================================================
    # PRIVATE HELPER METHODS
    # ========================================================================
    
    def _get_safe_value(self, key: str, default: str = '[MANUAL ENTRY REQUIRED]') -> str:
        """Safely extract value with fallback"""
        value = self.profile.get(key)
        
        if value is None:
            return default
        
        if isinstance(value, str):
            cleaned = sanitize_input(value.strip())
            return cleaned if cleaned else default
        
        if isinstance(value, list) and not value:
            return default
        
        return str(value)
    
    def _format_data_types_readable(self) -> str:
        """Convert coded data types to readable format"""
        types = self.profile.get('data_types', [])
        if not types:
            return '[NO DATA TYPES SPECIFIED]'
        
        readable = [
            DATA_TYPE_DISPLAY_NAMES.get(t, t.replace('_', ' ').title()) 
            for t in types
        ]
        return ', '.join(readable)
    
    def _format_data_types_detailed(self) -> str:
        """Format data types with detailed descriptions"""
        types = self.profile.get('data_types', [])
        if not types:
            return '[SPECIFY DATA TYPES]'
        
        formatted = []
        for dtype in types:
            desc = DATA_TYPE_DETAILED_DESCRIPTIONS.get(
                dtype, 
                dtype.replace('_', ' ').title()
            )
            formatted.append(f"- {desc}")
        
        return '\n'.join(formatted)
    
    def _format_purposes(self) -> str:
        """Format data processing purposes"""
        purposes = self.profile.get('purposes', [])
        
        if purposes:
            return ', '.join(purposes)
        
        # Generate defaults based on entity type
        entity = self.profile.get('entity_type', 'general')
        
        default_purposes = {
            'ecommerce': 'Order processing, payment processing, delivery, customer support',
            'social_media': 'Account management, content personalization, connections',
            'gaming': 'Game functionality, leaderboards, in-game purchases, matchmaking',
            'fintech': 'Financial transactions, credit assessment, fraud prevention, compliance',
            'healthcare': 'Medical treatment, health records, appointments, consultations',
            'edtech': 'Course delivery, progress tracking, certification, personalized learning'
        }
        
        default = default_purposes.get(
            entity, 
            'Service delivery, account management, customer support'
        )
        return f"{default} [REVIEW AND CUSTOMIZE]"
    
    def _format_withdrawal_method(self) -> str:
        """Generate consent withdrawal instructions"""
        has_consent = self.profile.get('has_consent_mechanism', False)
        
        if has_consent:
            return """You may withdraw consent at any time through:

1. Account Settings: Privacy Settings > Withdraw Consent
2. Email: [INSERT PRIVACY EMAIL] with subject "Withdraw Consent"
3. Phone: [INSERT PHONE] during business hours

Withdrawal processed within 48 hours. Service delivery may be affected if withdrawn 
data is necessary for contract performance."""
        else:
            return """[CRITICAL: Rule 3(4) Violation - No Withdrawal Mechanism]

Required Implementation (choose one or more):
- Account Settings toggle (recommended)
- Dedicated email address
- Toll-free phone number
- Physical mail with postage-paid envelope

Withdrawal must be as easy as giving consent.
Penalty: Rs. 50 crore"""
    
    def _format_retention_summary(self) -> str:
        """Summarize retention periods"""
        types = self.profile.get('data_types', [])
        if not types:
            return '[DEFINE RETENTION PERIODS]'
        
        max_retention = None
        for dtype in types:
            period = self._calculate_retention_period(dtype)
            if '10 years' in period:
                max_retention = '10 years'
            elif '7 years' in period and max_retention != '10 years':
                max_retention = '7 years'
            elif '5 years' in period and max_retention not in ['10 years', '7 years']:
                max_retention = '5 years'
        
        if max_retention:
            return f"Up to {max_retention} depending on data category (see Retention Schedule)"
        
        return "Until purpose completion or consent withdrawal (see Retention Schedule)"
    
    def _calculate_retention_period(self, data_type: str) -> str:
        """Calculate retention period based on legal obligations"""
        entity = self.profile.get('entity_type')
        user_count = self.profile.get('user_count', 0)
        
        # Third Schedule requirements
        if entity in ['ecommerce', 'social_media']:
            threshold = THIRD_SCHEDULE_THRESHOLDS.get(entity, float('inf'))
            if user_count >= threshold:
                if data_type in ['payment_info', 'name', 'email']:
                    return '3 years from account closure (Third Schedule)'
        
        if entity == 'gaming':
            threshold = THIRD_SCHEDULE_THRESHOLDS.get(entity, float('inf'))
            if user_count >= threshold:
                if data_type in ['name', 'email', 'behavioral']:
                    return '3 years from last activity (Third Schedule)'
        
        # Specific legal obligations
        if data_type == 'payment_info':
            return RETENTION_REQUIREMENTS['payment_info']
        if data_type == 'health_data':
            return RETENTION_REQUIREMENTS['health_data']
        if data_type == 'employment':
            return RETENTION_REQUIREMENTS['employment']
        
        return RETENTION_REQUIREMENTS['default']
    
    def _get_legal_basis(self, data_type: str) -> str:
        """Determine legal basis for processing"""
        if data_type == 'health_data':
            return LEGAL_BASIS_MAP['consent']
        
        if data_type == 'payment_info':
            entity = self.profile.get('entity_type')
            if entity in ['ecommerce', 'fintech']:
                return f"{LEGAL_BASIS_MAP['contract']} + {LEGAL_BASIS_MAP['legal_obligation']}"
            return LEGAL_BASIS_MAP['contract']
        
        return LEGAL_BASIS_MAP['consent']
    
    def _replace_all_placeholders(self, doc: Document, replacements: Dict[str, str]):
        """Replace placeholders throughout document"""
        # Paragraphs
        for paragraph in doc.paragraphs:
            self._replace_in_paragraph(paragraph, replacements)
        
        # Tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        self._replace_in_paragraph(paragraph, replacements)
        
        # Headers (not footers - contains metadata)
        for section in doc.sections:
            for paragraph in section.header.paragraphs:
                self._replace_in_paragraph(paragraph, replacements)
    
    def _replace_in_paragraph(self, paragraph, replacements: Dict[str, str]):
        """Replace placeholders in paragraph preserving formatting"""
        full_text = ''.join(run.text for run in paragraph.runs)
        
        has_placeholder = any(p in full_text for p in replacements)
        
        if has_placeholder:
            for placeholder, value in replacements.items():
                full_text = full_text.replace(placeholder, value)
            
            for run in paragraph.runs:
                run.text = ''
            
            if paragraph.runs:
                paragraph.runs[0].text = full_text
            else:
                paragraph.add_run(full_text)
    
    def _populate_retention_table(self, doc: Document):
        """Find and populate retention schedule table in template"""
        # Find table with marker
        retention_table = None
        for table in doc.tables:
            if table.rows and table.rows[0].cells:
                # Check if this table has the marker in row 2
                if len(table.rows) > 1:
                    first_cell_text = table.rows[1].cells[0].text.strip()
                if '{{RETENTION_TABLE}}' in first_cell_text:
                    retention_table = table
                    break
    
        if not retention_table:
            raise DocumentGenerationError(
            "Retention schedule template missing {{RETENTION_TABLE}} marker in table"
        )
    
        # Clear the marker from row 2, cell 1
        retention_table.rows[1].cells[0].text = ''
    
        # Remove the marker row entirely (we'll add data rows fresh)
        # Note: Can't delete row easily in python-docx, so we'll just overwrite it
    
        # Get data types
        data_types = self.profile.get('data_types', [])
    
        if not data_types:
        # If no data types, add placeholder row
            row = retention_table.rows[1]
            row.cells[0].text = '[NO DATA TYPES SPECIFIED]'
            row.cells[1].text = 'N/A'
            row.cells[2].text = 'N/A'
            row.cells[3].text = 'N/A'
            return
    
    # Populate first data type in existing row 2
        first_type = data_types[0]
        row = retention_table.rows[1]
        row.cells[0].text = DATA_TYPE_DISPLAY_NAMES.get(first_type, first_type.title())
        row.cells[1].text = self._calculate_retention_period(first_type)
        row.cells[2].text = self._get_legal_basis(first_type)
        row.cells[3].text = '48-hour warning + user confirmation'
    
        # Add additional rows for remaining data types
        for data_type in data_types[1:]:
            row = retention_table.add_row()
            row.cells[0].text = DATA_TYPE_DISPLAY_NAMES.get(data_type, data_type.title())
            row.cells[1].text = self._calculate_retention_period(data_type)
            row.cells[2].text = self._get_legal_basis(data_type)
            row.cells[3].text = '48-hour warning + user confirmation'

    def _add_document_metadata(self, doc: Document, doc_name: str, rule_ref: str):
        """Add metadata footer"""
        section = doc.sections[0]
        footer = section.footer
        
        for paragraph in footer.paragraphs:
            paragraph.clear()
        
        metadata = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        metadata.text = f"""Generated by DPDPA Compliance Dashboard v1.0.2 | {datetime.now().strftime('%B %d, %Y at %I:%M %p')}
Document: {doc_name} ({rule_ref}) | Based on: DPDP Act 2023 + Rules 2025 (Nov 13, 2025)
Business: {self.profile.get('business_name', 'Unknown')} | Document ID: {secrets.token_hex(8)}

TEMPLATE ONLY - Legal review required | github.com/Tushar-9802/DPDPA"""
        
        metadata.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        for run in metadata.runs:
            run.font.size = Pt(7)
            run.font.color.rgb = RGBColor(128, 128, 128)
    
    def _add_template_disclaimer(self, doc: Document):
        """Add disclaimer at start"""
        disclaimer = doc.paragraphs[0].insert_paragraph_before()
        
        disclaimer.add_run('LEGAL TEMPLATE - REVIEW REQUIRED\n').bold = True
        disclaimer.add_run(
            'This document was auto-generated. It is a TEMPLATE requiring review by a '
            'qualified data protection lawyer before deployment. The developer assumes NO '
            'LIABILITY for errors, omissions, or legal consequences. Yellow-highlighted '
            'sections require manual completion.\n\n'
        )
        
        disclaimer.paragraph_format.space_after = Pt(12)
        
        for run in disclaimer.runs:
            run.font.size = Pt(9)
    
    def _add_breach_template_header(self, doc: Document, recipient: str):
        """Add breach template header"""
        header = doc.paragraphs[0].insert_paragraph_before()
        
        header.add_run(f'BREACH NOTIFICATION TEMPLATE - FOR FUTURE USE\n').bold = True
        header.add_run(f'Recipient: {recipient}\n\n').bold = True
        
        header.add_run('DO NOT FILE THIS NOW. ').bold = True
        header.add_run(
            'Pre-prepared template for crisis preparedness.\n\n'
            
            'WHEN TO USE: Data breach has occurred and been confirmed\n'
            'DEADLINE: Within 72 hours of discovery\n\n'
            
            'HOW TO USE:\n'
            '1. Review with legal counsel NOW (before breach)\n'
            '2. Store with incident response plan\n'
            '3. Fill bracketed sections [LIKE THIS] when breach occurs\n'
            '4. Legal counsel review completed notification\n'
            '5. Submit within 72 hours\n\n'
            
            'PENALTY FOR LATE NOTIFICATION: Rs. 200 crore\n\n'
        )
        
        header.paragraph_format.space_after = Pt(18)